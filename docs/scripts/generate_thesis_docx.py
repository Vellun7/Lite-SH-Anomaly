#!/usr/bin/env python3
"""
将毕业论文完整稿.md转换为docx文档
使用参考论文的样式和格式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re
import os

# 读取毕业论文内容
input_path = "/Users/lukachihanbao/GraduationProject/Lite-SH-Anomaly/docs/毕业论文完整稿.md"
output_path = "/Users/lukachihanbao/GraduationProject/Lite-SH-Anomaly/docs/毕业论文.docx"

with open(input_path, 'r', encoding='utf-8') as f:
    content = f.read()

# 创建新文档
doc = Document()

# 设置默认字体为宋体（中文）和Times New Roman（英文）
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(level=level)
    heading.clear()
    run = heading.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(16)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(14)
        run.bold = True
    return heading

def add_paragraph(doc, text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)  # 左缩进2字符
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_formula_paragraph(doc, text):
    """添加公式段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

def parse_markdown(content):
    """解析markdown内容并转换为docx元素"""
    lines = content.split('\n')
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # 跳过分隔线
        if line.strip() == '---':
            i += 1
            continue

        # 处理标题
        if line.startswith('# '):
            result.append(('title', line[2:]))
        elif line.startswith('## '):
            result.append(('h1', line[3:]))
        elif line.startswith('### '):
            result.append(('h2', line[4:]))
        elif line.startswith('#### '):
            result.append(('h3', line[5:]))

        # 处理正文段落
        elif line.strip() and not line.startswith('#'):
            # 检查是否有连续的非空行属于同一段落
            para_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].startswith('#') and not lines[j].startswith('**') and not lines[j].startswith('|'):
                para_lines.append(lines[j])
                j += 1
            result.append(('para', ' '.join(para_lines)))
            i = j - 1

        # 处理加粗文本块
        elif line.startswith('**') and '**' in line[2:]:
            text = line.strip()
            if text.startswith('**') and text.endswith('**') and text.count('**') == 2:
                result.append(('bold_para', text[2:-2]))
            else:
                result.append(('para', text))

        # 处理表格（简单处理）
        elif line.startswith('|'):
            # 收集整个表格
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith('|'):
                table_lines.append(lines[j])
                j += 1
            result.append(('table', table_lines))
            i = j - 1

        # 处理图片
        elif line.startswith('!['):
            match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
            if match:
                alt_text, url = match.groups()
                result.append(('image', alt_text, url))

        i += 1

    return result

def create_table_from_lines(doc, lines):
    """从markdown表格行创建docx表格"""
    # 解析表格结构
    rows_data = []
    for line in lines:
        if line.startswith('|') and line.endswith('|'):
            # 移除首尾的|，分割各列
            cells = [c.strip() for c in line[1:-1].split('|')]
            # 过滤掉分隔行（包含---）
            if not all(re.match(r'^[-:]+$', c.replace(' ', '')) for c in cells if c.strip()):
                rows_data.append(cells)

    if len(rows_data) < 2:
        return None

    # 创建表格
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            # 设置单元格字体
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)

    return table

# 解析内容
elements = parse_markdown(content)

# 生成文档
for elem in elements:
    if elem[0] == 'title':
        # 文档标题（居中大字）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(elem[1])
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(22)
        run.bold = True

    elif elem[0] == 'h1':
        # 一级标题
        add_heading(doc, elem[1], 1)

    elif elem[0] == 'h2':
        # 二级标题
        add_heading(doc, elem[1], 2)

    elif elem[0] == 'h3':
        # 三级标题
        add_heading(doc, elem[1], 3)

    elif elem[0] == 'para':
        text = elem[1].strip()
        if text:
            # 处理公式（$$...$$）
            if '$$' in text:
                parts = re.split(r'(\$\$[^$]+\$\$)', text)
                for part in parts:
                    if part.startswith('$$') and part.endswith('$$'):
                        add_formula_paragraph(doc, part[2:-2])
                    elif part.strip():
                        add_paragraph(doc, part)
            else:
                add_paragraph(doc, text)

    elif elem[0] == 'bold_para':
        text = elem[1].strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
            run.bold = True

    elif elem[0] == 'table':
        create_table_from_lines(doc, elem[1])

# 保存文档
doc.save(output_path)
print(f"文档已保存到: {output_path}")
print(f"文档包含 {len(doc.paragraphs)} 个段落")