#!/usr/bin/env python3
"""
以参考论文为样式模板，生成包含毕业论文内容的新文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

# 参考论文路径（用于获取样式）
reference_path = "/Users/lukachihanbao/Documents/刘燊林-基于深度神经网络的农业病虫害智能识别系统-终稿.docx"
# 毕业论文内容
thesis_path = "/Users/lukachihanbao/GraduationProject/Lite-SH-Anomaly/docs/毕业论文完整稿.md"
# 输出路径
output_path = "/Users/lukachihanbao/GraduationProject/Lite-SH-Anomaly/docs/毕业论文_完整版.docx"

# 读取参考论文以获取样式
ref_doc = Document(reference_path)

# 读取毕业论文内容
with open(thesis_path, 'r', encoding='utf-8') as f:
    thesis_content = f.read()

# 创建新文档
new_doc = Document()

# 复制页面边距
for section in ref_doc.sections:
    new_section = new_doc.sections[0] if new_doc.sections else new_doc.add_section()
    new_section.top_margin = section.top_margin
    new_section.bottom_margin = section.bottom_margin
    new_section.left_margin = section.left_margin
    new_section.right_margin = section.right_margin

# 设置默认字体
style = new_doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置标题样式
def apply_heading_style(run, level):
    """应用标题样式"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(16)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(14)
        run.bold = True
    elif level == 4:
        run.font.size = Pt(12)
        run.bold = True

def add_formatted_paragraph(doc, text, first_indent=True, bold=False):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_center_paragraph(doc, text, bold=False, size=12):
    """添加居中段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    return p

# 解析内容并生成文档
lines = thesis_content.split('\n')

# 状态跟踪
in_preface = False
in_introduction = False
in_ref = False
in_acknowledgements = False
current_h1 = None

i = 0
while i < len(lines):
    line = lines[i]

    # 跳过分割线
    if line.strip() == '---':
        i += 1
        continue

    # 主标题
    if line.startswith('# ') and not line.startswith('## '):
        title = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:].strip())

        if title == "面向智能家居的轻量化异常检测算法设计与实现":
            # 文档标题（居中）
            p = new_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            in_preface = True

        elif title == "设计总说明":
            p = new_doc.add_heading(level=1)
            run = p.add_run(title)
            apply_heading_style(run, 1)
            in_preface = True

        elif title == "Introduction":
            p = new_doc.add_heading(level=1)
            run = p.add_run("Abstract")
            apply_heading_style(run, 1)
            in_introduction = True

        elif title == "目  录":
            p = new_doc.add_heading(level=1)
            run = p.add_run(title)
            apply_heading_style(run, 1)

        elif "参考文献" in title:
            p = new_doc.add_heading(level=1)
            run = p.add_run(title)
            apply_heading_style(run, 1)
            in_ref = True
            in_preface = False
            in_introduction = False

        elif "致谢" in title:
            p = new_doc.add_heading(level=1)
            run = p.add_run(title)
            apply_heading_style(run, 1)
            in_acknowledgements = True
            in_ref = False

        elif title == "总结与展望":
            p = new_doc.add_heading(level=1)
            run = p.add_run(title)
            apply_heading_style(run, 1)
            in_acknowledgements = False
            in_preface = False

        else:
            # 普通章节
            p = new_doc.add_heading(level=1)
            run = p.add_run(title)
            apply_heading_style(run, 1)
            in_preface = False
            in_introduction = False
            current_h1 = title

    # 二级标题
    elif line.startswith('## '):
        title = line[3:].strip()
        p = new_doc.add_heading(level=2)
        run = p.add_run(title)
        apply_heading_style(run, 2)

    # 三级标题
    elif line.startswith('### '):
        title = line[4:].strip()
        p = new_doc.add_heading(level=3)
        run = p.add_run(title)
        apply_heading_style(run, 3)

    # 四级标题
    elif line.startswith('#### '):
        title = line[5:].strip()
        p = new_doc.add_heading(level=4)
        run = p.add_run(title)
        apply_heading_style(run, 4)

    # 加粗段落（关键词等）
    elif line.strip().startswith('**') and line.strip().endswith('**'):
        text = line.strip()[2:-2]
        if '关键词' in text or 'Key Words' in text:
            p = new_doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        else:
            add_formatted_paragraph(new_doc, text, bold=True)

    # 表格
    elif line.startswith('|'):
        if '---' not in line:
            p = new_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clean_text = line.replace('|', '  ')
            run = p.add_run(clean_text)
            run.font.size = Pt(10)

    # 正文段落
    elif line.strip():
        if in_ref and line.strip().startswith('['):
            # 参考文献格式
            p = new_doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.left_indent = Cm(0.74)
            run = p.add_run(line.strip())
            run.font.size = Pt(10.5)
        elif in_preface or in_introduction or in_acknowledgements:
            # 前言部分保持原样
            add_formatted_paragraph(new_doc, line.strip())
        else:
            # 检查是否有公式
            if '$$' in line:
                parts = re.split(r'(\$\$[^$]+\$\$)', line)
                for part in parts:
                    if part.startswith('$$') and part.endswith('$$'):
                        p = new_doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(part[2:-2])
                        run.font.size = Pt(11)
                    elif part.strip():
                        add_formatted_paragraph(new_doc, part)
            else:
                add_formatted_paragraph(new_doc, line.strip())

    i += 1

# 保存文档
new_doc.save(output_path)
print(f"文档已保存到: {output_path}")
print(f"文档包含 {len(new_doc.paragraphs)} 个段落")

# 验证输出
verify_doc = Document(output_path)
print(f"\n=== 文档前60个段落 ===")
for i, p in enumerate(verify_doc.paragraphs[:60]):
    text = p.text[:65] if p.text else "[空]"
    style = p.style.name if p.style else "None"
    if text.strip():
        print(f"[{i:3d}] {style[:15]:15s} | {text}")